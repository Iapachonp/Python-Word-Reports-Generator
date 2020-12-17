from tkinter import *
from tkinter.ttk import *
from tkinter import filedialog 
from tkinter import scrolledtext
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.text.run import Font, Run
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING
import warnings                                  # modo `no molestar`
warnings.filterwarnings('ignore')
#! pip install python-docx    use and create word documents in python! 

filename= str()


##----------------------------- Functions Assessments -------------------------------------------------

def Load_Assessment_DB(Assess_Name , Assess_Lang):
    
    Url = r".\Web form Data bases\TXT data bases\\"
    Url = Url + Assess_Name + "\\" + Assess_Lang
    
    
    with open(Url+"\Answers.txt", "r",  encoding="utf-8") as file:
        C = eval(file.readline())
    with open(Url+"\Questions.txt", "r",  encoding="utf-8") as file:
        A = eval(file.readline())
    
    #for i in range(0, len(A)):
        
        #print(str(i)+'.'+' '+ A[i]+ str(C[i]))
        
    right_ans= []

    for i in range(3, len(C)):
        ans= list(C[i].keys())
        for j in range(0, len(ans)):
            get_ans_val= C[i].get(ans[j])
            #print(get_ans_val)
            if str(get_ans_val) == "{True}":
                right_ans.append(str(ans[j]))
                #print("yuju")

    return (A,C, right_ans)

def Load_Participants_Assess(url, Exam, lang):
    
    Asses = pd.read_excel (url)
    
    Asses.drop(['Quiz feedback' , 'Name', 'Email', 'Points - Name',
       'Feedback - Name', 'Points - Email', 'Feedback - Email', 'Points - Distributor', 'Feedback - Distributor'], axis=1 ,inplace = True)
    
    Total_Points = (len(Asses.columns)-7)/3
    Grade = []
    Wrong = []
    Duration= []
    Id= []
    
    for j in range (0, Asses.shape[0]):
        Grade.append(round(Asses.iloc[j]['Total points']/Total_Points*100))
        Wrong.append(Total_Points - Asses.iloc[j]['Total points']) 
        Duration.append(str(Asses.iloc[j]['Completion time']- Asses.iloc[j]['Start time'])[8:]+ (" HH/mm/ss"))
        Id.append(Exam + "-" + lang+ "-" + str(Asses.iloc[j]['ID']))
    
    Asses['Grade'] = Grade
    Asses['Wrong answers'] = Wrong
    Asses['Completion time'] = Duration
    Asses['ID'] = Id
    
    Asses.drop(['Start time'], axis=1, inplace= True)
                               
    #Change Grade, Duration and Wrong columns position
    
    cols = list(Asses) # get a list of columns
    cols.insert(3, cols.pop(cols.index('Wrong answers')))  # move the column to head of list using index, pop and insert
    Asses= Asses.loc[:, cols]  # use ix to reorder
    
    cols = list(Asses) # get a list of columns
    cols.insert(2, cols.pop(cols.index('Grade')))  # move the column to head of list using index, pop and insert
    Asses= Asses.loc[:, cols]  # use ix to reorder
    
    # Rename some columns to have a correct name in the report
    Asses.rename(columns = {'Name2':'Name'}, inplace = True) 
    Asses.rename(columns = {'Email2':'Email'}, inplace = True) 

    
    return (Asses)

def Individual_Results(Asses1,ques, ans ,  Correct_ans, Latam_ave, Number_asess):
    
    
    question_flag = 0

    
    
    #Create first page 
    #section = document.sections[0]
    #document.add_heading('' , 1)  # add a Space header
    #document.add_picture(r"C:\Users\Ipachon1\OneDrive - Rockwell Automation, Inc\LATAM\ASSESMENTS\others\RA-Partner.png", width=Inches(5.8))
            
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    #document.add_heading('' , 0)  # add a Space header
    
    
    #document.add_heading('Distributor Competency Specialist Assessment Results - 2020 ' , 0)  # add a header
    #First_paragraph = document.add_paragraph(style='Body Text')  
    #First_paragraph.add_run('This report includes the competency test results, we encourage you to study this information and consult the documents, links and articles suggested so you can improve your proficiency.')
    #document.add_page_break() 
    
    
    
    
    for j in range (0, Asses1.shape[0]): #trough all the assessments of the Data base

        first = 0

        #Create the Word Document
        document = Document()  # Create it

        #footer creation   
        section = document.sections[0]
        Footer = section.footer
        paragraph = Footer.paragraphs[0]
        paragraph.text = "Rockwell Automation \t\tStrictly Confidential information Do not distribute"
        
        document.add_heading(" Assessment Results for: " + str(Asses1.iloc[j].Name) , 1)  # add a header
        Second_paragraph = document.add_paragraph(style='Body Text')  
        Grade_Comparation = 'Higher' if Latam_ave < Asses1.iloc[j].Grade else 'Lower'  #Grade comparation string to include in the paragraph
        Second_paragraph.add_run(str(Asses1.iloc[j].Name)+'’s total grade in this assessment was '+ str(Asses1.iloc[j].Grade)+ '% which is ' + Grade_Comparation + ' than the general average taken over ' + str(Number_asess) +' specialists around Latin America. Overall average for Latin America is ' + str(Latam_ave)+ '%')   # add a information paragraph 
        
        
        document.add_heading('General Information', level=1) #add general information header 
        fourth_paragraph = document.add_paragraph(style= 'Body Text') # add a paragraph 
        fourth_paragraph.add_run('This section includes specialist information and the overall results of the 2020 assessment.')
        Space_paragraph = document.add_paragraph(style='Body Text')  # add a space as a paragraph
        Space_paragraph.add_run('')
        
        for i in range (0, Asses1.shape[1]-1):  #trough all the columns of the Data base
                if Asses1.columns[i].startswith( "Points") or Asses1.columns[i].startswith( "Feedback"):  #skip feed-back and points columns
                    pass
                else:
                    
                    if len(Asses1.columns[i]) > 20:             #this is a quick fix in order to correct only real questions
                        question_flag = question_flag +1
                        if Asses1.iloc[j,i+1] == 1:  #see if the answer is correct or not
                            #Answer.add_run("Your Response--"+str(Asses1.iloc[j,i])).font.color.rgb = RGBColor(0, 153, 0) #add answer green- correct answer
                            pass
                        else: 
                            
                            
            ########################## first time the questions start to be printed #####################################
                            
                            if first == 0: 
                                print("fuckkkkkkk"+str(i))
                                document.add_page_break()   #add a page break when questions are about to be displayed
                                document.add_heading(" Assessment Results for: " + str(Asses1.iloc[j].Name) , 1)  # add a header
                                Third_paragraph = document.add_paragraph(style='Body Text') # add a information paragraph
                                Third_paragraph.add_run('The following report provides a detailed list of results by APR and by specialist. For each question that was not correctly answered we provide below the question, the list of options, the correct answer, and the sources to learn more about this specific subject.')  # add a information paragraph
                                document.add_heading('' , 0)  # add a Space header
                            
                            first = first +1    
                            
            ########################## first time the questions start to be printed #####################################
                            
            
            
            ################change this in order to display all answers not only wrong answers#####################
                            Question= document.add_paragraph(style='Body Text')  #create Question Paragraph
                            Question.add_run(str(question_flag)+ (") ")+str(ques[3+question_flag-1])).bold = True   #add normal questions
                    
                            
                            All_ans= list(ans[3+question_flag-1].keys())
                            for k in range(0, len(All_ans)):
                                Question= document.add_paragraph(style='List Bullet')  #create Question Paragraph
                                Question.add_run(str(All_ans[k])) #add normal questions
                                
                                
                                
                            
                            Answer= document.add_paragraph(style='List Bullet')    #create Answers Paragraph               
           ################change this in order to display all answers not only wrong answers#####################   
                            
                            Answer.add_run("Your Response--"+str(Asses1.iloc[j,i])).font.color.rgb = RGBColor(168, 0, 0) #add answer red- Incorrect answer
                            Correct_Answer= document.add_paragraph(style='List Bullet') # Create Correct Answer paragraph 
                            Correct_Answer.add_run(" ---- Correct Answer-----" + str(Correct_ans[question_flag-1])).italic = True  #add correct answer from the database

                            Learning_Data= document.add_paragraph(style='List Bullet') # Create Correct Answer paragraph 
                            Learning_Data.add_run(" ----Reference Materials -----" + str(Asses1.iloc[j,i+2])).italic = True  #add learning data from the database
                            
                            
                    else:
                        
            ################change this in order to display all answers not only wrong answers#####################
                        Question= document.add_paragraph(style='Body Text')  #create Question Paragraph
                        Question.add_run(str(Asses1.columns[i])).bold = True   #add normal questions
                        Answer= document.add_paragraph(style='List Bullet')    #create Answers Paragraph
            ################change this in order to display all answers not only wrong answers#####################           
                        
                        Answer.add_run(str(Asses1.iloc[j,i]))  # add questions that are not from the assessment (name, email, etc)

        #document.add_page_break()
        print(j)
        document.save(r".\Results\\" +  str(Asses1.iloc[j].Distributor)+ "--Email--" + str(Asses1.iloc[j].Email)+ "--User Name--" + str(Asses1.iloc[j].Name)+ "--ID--" + str(Asses1.iloc[j].ID)+'.docx')  #Save Document;  Name 
        
        #first = 0
        question_flag = 0
        
    
    return str("The docx have been created, check the RESULTS folder")

##----------------------------- Functions TKINTER -------------------------------------------------

def clicked():
    
    Ques, Ans, Right_ans = Load_Assessment_DB(str(combo.get()), str(selected.get())) #English (Original)  Portuguese
    Asss= Load_Participants_Assess(filename, str(combo.get()), str(selected.get()))
    popupmsg(str(Individual_Results(Asss, Ques, Ans, Right_ans, 55.47, 28)))



# Function for opening the  
# file explorer window 
def browseFiles(): 
    global filename
    filename = filedialog.askopenfilename(initialdir = "/", 
                                          title = "Select a File", 
                                          filetypes = (("Excel Files", 
                                                        "*.xlsx*"), 
                                                       ("all files", 
                                                        "*.*"))) 
       
    # Change label contents 
    label_file_explorer.configure(text="Selected file : "+filename, font=("Arial Bold", 10)) 

def popupmsg(msg):
    popup = Tk()
    popup.wm_title("Successful Job Done!")
    popup.geometry('350x200')
    txt = scrolledtext.ScrolledText(popup,width=40,height=10, pady=40)
    txt.insert(INSERT,msg)
    txt.grid(column=0,row=1)
    #label = Label(popup, text=msg, font=("Arial Bold", 10))
    #label.pack(side="top", fill="x", pady=40)
    B1 = Button(popup, text="Great", command = popup.destroy)
    B1.pack()
    popup.mainloop()

##----------------------------- Functions  -------------------------------------------------


##----------------------------- GUI Design  -------------------------------------------------

window = Tk()

window.title("Market Access")
window.geometry('350x350')

lbl = Label(window, text="Results Generator" , font=("Arial Bold", 20))
lbl.grid(column=1, row=0)


combo_lbl = Label(window, text="Select the Assessment type" , font=("Arial Bold", 10))
combo_lbl.grid(column=1, row=4)
combo = Combobox(window)
combo['values']= ('CSM', 'IC', 'IE', 'IS', 'Logix','LV', 'LV MCC', 'LV MV MCC', 'Motion', 'MV', 'Visualization')
combo.current(0) #set the selected item
combo.grid(column=1, row=6)


radio_lbl = Label(window, text="Select the Assessment language" , font=("Arial Bold", 10))

selected = StringVar()
rad1 = Radiobutton(window,text='English', value='English', variable=selected)
rad2 = Radiobutton(window,text='Spanish', value='Spanish', variable=selected)
rad3 = Radiobutton(window,text='Portuguese', value='Portuguese', variable=selected)




btn = Button(window, text="Generate Results", command=clicked , width=25)



# Create a File Explorer label 
label_file_explorer = Label(window,  
                            text = "Select the Results Excel file from forms:", 
                            width = 50, font=("Arial Bold", 10)
                            ) 
   
       
button_explore = Button(window,  
                        text = "Browse Files", 
                        command = browseFiles)  
   
button_exit = Button(window,  
                     text = "Exit", 
                     command = window.destroy)  
   
# Grid method is chosen for placing 
# the widgets at respective positions  
# in a table like structure by 
# specifying rows and columns 
label_file_explorer.grid(column = 1, row = 1) 
radio_lbl.grid(column=1, row=7)
rad1.grid(column=1, row=8)
rad2.grid(column=1, row=9)
rad3.grid(column=1, row=10)
   
button_explore.grid(column = 1, row = 2) 
   
btn.grid(column=1, row=14, pady=10)
button_exit.grid(column = 1,row = 15) 


window.iconbitmap(r'.\others\Pictures\logo.ico')
window.mainloop()


##----------------------------- GUI Design  -------------------------------------------------
